
#!/usr/bin/env python38
#!c:/Python38/python38.exe
# -*- coding: utf-8 -*-

########################################################################################################################
### Module   = CopyCtr_GenDocFile.py                                                                                 ###
### Author   = Truong Pham                                                                                           ###
### Date     = 12/29/2021                                                                                            ###
###                                               Revision History                                                   ###
###                                                                                                                  ###
###  1.0.0: 12/29/2021  : Initial version.                                                                           ###
########################################################################################################################
###                                                Import section                                                    ###
########################################################################################################################


import os
import re
from shutil import copy2
import sys
import argparse
from os import environ, makedirs
import docx
from docx.shared import Pt
import glob
import pycicd
import json
import shutil
import subprocess
from pathlib import Path
from Global_Properties import Global_Properties
from C0_Measure import C0_Measure   
from Correct_Ctr import Correct_Ctr_File 
########################################################################################################################
###                                            Global Variables deffition                                            ###
########################################################################################################################

glbWorkSpace = os.environ['WORKSPACE'].replace('\\', '/')   
# glbWorkSpace = ''
glbCantataWS = ''
glbMergeWS = ''
indent = "  "
jobProperties = Global_Properties()
print(jobProperties)
LOGGER = pycicd.ci_logger('Copy Result').logger
########################################################################################################################
###                                                Function definition                                               ###
########################################################################################################################


def Create_DocFile(Information):
    for Filename in Information:
        file = glob.glob( f'{glbCantataWS}/*/src/{Filename}.c')
        # Function check content of C files - T.B.D
        file = file[0]
        sourcefile = open(file, 'r')
        content = sourcefile.read()
        my_doc = docx.Document()
        para = my_doc.add_paragraph().add_run(content)
        para.font.size = Pt(9)
        para.font.name = 'Courier New'
        os.makedirs(f'{glbMergeWS}/input',exist_ok=True)
        os.makedirs(f'{glbMergeWS}/log',exist_ok=True)
        my_doc.save(
            f'{glbMergeWS}/input/{Filename}_c.doc')
        #print(f'>>> Created {glbMergeWS}/input/{prj}/{Filename}_c.doc\n')
        if not os.path.exists(f'{glbMergeWS}/result'):
            makedirs(
                f'{glbMergeWS}/result')
        copy2(f'{glbMergeWS}/input/{Filename}_c.doc',
              f'{glbMergeWS}/result/{Filename}_c_result.doc')
        #print(f'>>> Copied {glbMergeWS}/input/{prj}/{Filename}_c.doc to {glbMergeWS}/result/{Filename}_c_result.doc\n')
        code = open(
            f'{glbMergeWS}/log/{Filename}_c_code.txt', 'w')
        code.write(content)
        code.close()
        open(
            f'{glbMergeWS}/log/{Filename}_c_notcomp.txt', 'a').close()
        LOGGER.success(f'Generated doc file for {Filename}.c')


def Copy_cpl_cpr():
    """
    This method will copy cpl and ctr files To report path and correct ctr content also.
    """
    cplFiles = glob.glob(
        f'{glbCantataWS}/*/Cantata/tests/**/*.cpl', recursive=True)
    ctrFiles = glob.glob(
        f'{glbCantataWS}/*/Cantata/tests/atest_*/*.ctr', recursive=True)
    projectFomart = re.compile(f'{glbCantataWS}/(.*)/Cantata/.*')
    # copy cpl files
    for cplPath in cplFiles:
        cplPath = cplPath.replace("\\", "/")
        projectName = projectFomart.search(cplPath).group(1)

        copy2(cplPath,
              f'{glbMergeWS}/input/{os.path.basename(cplPath)}') #Author Huy doan: Fix override file cpl
        LOGGER.success(f'Copied {os.path.basename(cplPath)}')
    # copy ctr files and correct it.
    for ctrPath in ctrFiles:
        ctrPath = ctrPath.replace("\\", "/")
        projectName = projectFomart.search(ctrPath).group(1)
        copy2(ctrPath,
              f'{glbMergeWS}/input/{os.path.basename(ctrPath)}')#Author Huy doan: Fix override file cpr
        # Update to support issue for ctr wrong name.
        Correct_Ctr_File().repairCtrFiles(f'{glbMergeWS}/input')#Author Huy doan: Fixing to repair cpr
        LOGGER.success(f'Copied {os.path.basename(ctrPath)}')
        

def get_file_list():
    """
    return: [list of c file name only - which is compiled successful]
    """
    allObjectFiles = glob.glob(f"{glbCantataWS}/*/Cantata/tests/atest_*/src/*.o", recursive=False)
    allObjectFiles = [Path(x).stem for x in allObjectFiles]
    allObjectFiles = list(set(allObjectFiles))
    with open(f'{glbMergeWS}/target.txt', 'w', newline='\n') as f:
        # should add \n at the last row to solve CMT.sh Workspaceissue
        f.write("\n".join(allObjectFiles) + '\n')
    return allObjectFiles

def CopyResult():
    jsonData = jobProperties.targetCantataDir
    with open(jsonData, 'r') as f:
        data = json.load(f)
    
    global glbCantataWS,glbMergeWS
    glbCantataWS = 'C:/Workspace/CantataWorkSpace'
    rootReport = 'C:/Workspace/CantataWorkSpace/S4_CR52_EthSwtCont_TestLog/Env'
    glbMergeWS = 'C:/Workspace/Create_DUT_Report/rha_cicd-master-ci/02_Jobs/50_UT/03_DUT/scripts/Create_UnitTest_Report/CovReportMerge/Env/msn'
    templateReport = f'{rootReport}/msn'
    shutil.copytree(templateReport, glbMergeWS, dirs_exist_ok=True)
    Information = get_file_list()
    LOGGER.info(str(Information))
    Create_DocFile(Information)
    Copy_cpl_cpr()

def generateLogResult():
    ########################################################################################################################
    ###                                                MAIN FUNCTION                                                     ###
    ########################################################################################################################
    # Sequence called
    # Get user input : This path the same as target.txt path
    baseDir = glbMergeWS  # "D:/sendToVendor/CantataWS/CantataWS/X2x/U2Ax/gpt"
    baseDir = baseDir.replace("\\", "/")
    print("The path of base dir" , baseDir)
    # Init C0 class
    C0_Init = C0_Measure(baseDir, baseDir)
    # Create Cpr File All_preprocessorMacros.cpr
    C0_Init.create_CprFile()
    # Filter to get un-complied macros via C
    unCompiledPerSource = C0_Init.getLoCUnCompiledMacros()
    # update un-compiled line of codes to *_notcomp.txt in log path of CMT
    logPath = "{}/log".format(baseDir)
    for source, locList in unCompiledPerSource.items():
        print(source, locList)
        linePerSource = C0_Init.listUnCompiledLineOfCode(logPath, source, locList)
    # re-confirm stage to check any errors
    if [] == C0_Init.errorLog:
        print("Update *_notcomp.txt successfull!")
    else:
        errorLogContent = "\n".join(C0_Init.errorLog)
        # print to jenkin log
        print("Error: Please re-check error log\n", errorLogContent)
        errorLogFile = "{}/error_C0_Measure.log".format(baseDir)
        # Print to error log file
        with open(errorLogFile, 'w') as log:
            log.write(errorLogContent)
            log.close()

def generateCtrResult():
    """
    Just call CMT.sh file to generate related files for reporting.
    """
    global glbMergeWS
    # call CMT script
    print(f"JK_INFO: Target Workspace {glbMergeWS}")
    subprocess.call(f"sh.exe ./CMT_execute.sh", cwd=glbMergeWS, timeout=10000, stdout=True)
    return 0  

def main():
    """Main function"""
    global glbCantataWS,glbMergeWS
    CopyResult()
    generateCtrResult()
    generateLogResult()

if __name__ == "__main__":
    main()

