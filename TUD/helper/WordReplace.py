import re
from docx import Document
from docx.text.run import Run
from docx.shared import RGBColor
document = Document("U:/internal/Module/memacc/06_CD/01_WorkProduct/generator_cs/doc/output_dd.docx")

Common = list(range(1,6 +1))
U5LxCommon = list(range(7, 11 +1))
U5L1 = list(range(12, 16+1))
U5L2 = list(range(17, 21+1))
U5L4 = list(range(22, 26+1))
# U2xCommon = list(range(36, 40 +1))
# U2A16 = list(range(41, 45 +1))
# U2A8 = list(range(51, 55 +1))
# U2A6 = list(range(46, 50 +1))
# U2BxCommon = list(range(36, 40 +1))
# U2B24 = list(range(41, 45 +1))
# U2B20 = list(range(51, 55 +1))
# U2B10 = list(range(46, 50 +1))
# U2B6 = list(range(46, 50 +1))
# U2CxCommon = list(range(95, 101 +1))
# U2C8 = list(range(117, 121 +1))
# U2C6 = list(range(112, 116 +1))
# U2C4 = list(range(107, 111 +1))
# U2C2 = list(range(102, 106 +1))

for p in document.paragraphs:
    rs = p._element.xpath('.//w:t')
    for run in rs:
        text = run.text
        match = re.search("MEMACC_TUD_CLS_[0-9]{3}_[0-9]{3}:",text)
        if match:
            number = int(text.split("_")[3])
            line = p.add_run()
            line.add_break()
            r = p.add_run()
            text1 = str("Dev: U5L4, U5L2, U5L1")
            if number in Common:
                text2 = str("[Attr: U5L4, U5L2, U5L1]")
            elif number in U5LxCommon:
                text2 = str("[Attr: U5L4, U5L2, U5L1]")
            elif number in U5L1:
                text2 = str("[Attr: -, -, U5L1]")
            elif number in U5L2:
                text2 = str("[Attr: -, U5L2, -]")
            elif number in U5L4:
                text2 = str("[Attr: U5L4, -, -]")
            # # elif number in U2xCommon:
            #     # text2 = str("[Attr: -, -, -, U2A16, U2A8]")
            # # elif number in U2A16:
            #     # text2 = str("[Attr: -, -, -, U2A16, -]")
            # # elif number in U2A8:
            #     # text2 = str("[Attr: -, -, -, -, U2A8]")
            # # elif number in U2A6:
            #     # text2 = str("[Attr: -, -, -, -, U2A6]")    
            # elif number in U2CxCommon:
            #     text2 = str("[Attr: -, -, -, -, -, -, -, -, -, -, U2C8, U2C6, U2C4, U2C2]")
            # elif number in U2C8:
            #     text2 = str("[Attr: -, -, -, -, -, -, -, -, -, -, U2C8, -, -, -]")
            # elif number in U2C6:
            #     text2 = str("[Attr: -, -, -, -, -, -, -, -, -, -, -, U2C6, -, -]")
            # elif number in U2C4:
            #     text2 = str("[Attr: -, -, -, -, -, -, -, -, -, -, -, -, U2C4, -]")
            # elif number in U2C2:
            #     text2 = str("[Attr: -, -, -, -, -, -, -, -, -, -, -, -, -, U2C2]")                
            r.text = text1
            r.bold = False
            r.add_break()
            r2 = p.add_run()
            r2.text = text2
            font = r2.font
            font.color.rgb = RGBColor(0x00, 0xBF, 0xFF)
            r2.bold= False
            r2.italic = True
            r2.add_break()

document.save('new3.docx')
